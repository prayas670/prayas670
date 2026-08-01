"""
Core analysis engine for the resume analyzer. Runs entirely locally with
regex/scikit-learn, so the app works with zero API keys.
"""

import os
import re
import difflib
from collections import Counter
from datetime import datetime

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from flashtext import KeywordProcessor

import ml_models
from ml_models import (
    sbert_similarity,
    extract_entities_spacy,
    build_ats_feature_vector,
    predict_ats_score,
)

def get_st_model():
    # Backwards-compatible alias; real implementation lives in ml_models.
    return ml_models.get_sbert_model()

SKILL_DB = {
    "languages": [
        "python", "java", "javascript", "typescript", "c++", "c#", "go", "golang",
        "rust", "ruby", "php", "swift", "kotlin", "scala", "r", "matlab", "sql",
        "html", "css", "bash", "shell", "perl", "objective-c", "dart",
    ],
    "frameworks": [
        "react", "angular", "vue", "next.js", "django", "flask", "fastapi",
        "spring", "spring boot", "express", "node.js", "rails", ".net", "laravel",
        "tensorflow", "pytorch", "keras", "scikit-learn", "pandas", "numpy",
        "bootstrap", "tailwind", "jquery", "redux", "graphql",
    ],
    "tools": [
        "git", "docker", "kubernetes", "jenkins", "aws", "azure", "gcp",
        "terraform", "ansible", "linux", "jira", "confluence", "figma",
        "postman", "webpack", "ci/cd", "github actions", "nginx", "redis",
        "mongodb", "postgresql", "mysql", "elasticsearch", "kafka", "spark",
        "hadoop", "tableau", "power bi", "excel", "airflow", "grafana",
    ],
    "soft_skills": [
        "leadership", "communication", "teamwork", "problem solving",
        "project management", "time management", "collaboration",
        "critical thinking", "adaptability", "mentoring", "public speaking",
        "negotiation", "agile", "scrum", "stakeholder management",
    ],
}
ALL_SKILLS = sorted({s for group in SKILL_DB.values() for s in group}, key=len, reverse=True)

ACTION_VERBS = {
    "achieved", "accelerated", "built", "created", "delivered", "designed",
    "developed", "drove", "engineered", "established", "executed", "expanded",
    "generated", "improved", "implemented", "increased", "initiated", "launched",
    "led", "managed", "optimized", "orchestrated", "pioneered", "reduced",
    "resolved", "spearheaded", "streamlined", "transformed", "boosted",
    "automated", "architected", "negotiated", "mentored", "scaled",
    "deployed", "served", "shipped", "maintained", "integrated", "authored",
    "migrated", "configured", "monitored", "processed", "analyzed", "tested",
    "debugged", "refactored", "wrote", "trained", "containerized",
    "provisioned", "presented", "published", "collaborated", "coordinated",
    "supported", "enhanced", "upgraded", "consolidated", "standardized",
    "modernized", "documented", "administered", "diagnosed", "instrumented",
}

WEAK_PHRASES = [
    "hardworking", "team player", "results-driven", "results driven",
    "detail-oriented", "detail oriented", "go-getter", "self-starter",
    "think outside the box", "synergy", "dynamic individual", "people person",
    "hard worker", "responsible for", "duties included",
]

VERB_HINTS = [
    (r"\b(manag|oversaw|oversee|supervis)\w*\b", "Managed"),
    (r"\b(built|build|develop|creat)\w*\b", "Built"),
    (r"\b(design)\w*\b", "Designed"),
    (r"\b(improv|optimi[sz]e|enhanc)\w*\b", "Optimized"),
    (r"\b(lead|led)\w*\b", "Led"),
    (r"\b(analy[sz]e|research)\w*\b", "Analyzed"),
    (r"\b(automat)\w*\b", "Automated"),
    (r"\b(reduc|decreas|cut|lower)\w*\b", "Reduced"),
    (r"\b(increas|grow|grew|boost|scal)\w*\b", "Increased"),
    (r"\b(launch|deploy|releas|ship)\w*\b", "Launched"),
    (r"\b(coordinat|organiz)\w*\b", "Coordinated"),
    (r"\b(support|assist|help)\w*\b", "Supported"),
]
DEFAULT_VERB = "Drove"

WEAK_OPENERS_RE = re.compile(
    r"^(responsible for|duties included|worked on|helped with|helped to|"
    r"assisted with|tasked with|in charge of)\s*",
    re.I,
)


# Matches real-world header variants ("Key Projects", "3. Projects", etc.)
# by allowing a leading qualifier word and decorative bullets/dashes, while
# still requiring the header be essentially standalone.
_HEAD_LEAD = r"^\s*[-=~#>*•\u2022\u2500\u2501\u2504\u2508│┃]*\s*(\d+[\.\)]\s*)?"
_HEAD_TAIL = r"\s*[-=~#>*•\u2022\u2500\u2501\u2504\u2508│┃:|]*\s*$"

SECTION_HEADER_PATTERNS = [
    ("summary", _HEAD_LEAD + r"(professional\s+|career\s+)?(summary|objective|profile|about\s*me)" + _HEAD_TAIL),
    ("experience", _HEAD_LEAD + r"((professional|relevant|work)\s+)?(experience|employment(\s+history)?|work\s+history|internships?)" + _HEAD_TAIL),
    ("education", _HEAD_LEAD + r"(education(al)?(\s+background)?|academics?|academic\s+background|qualifications)" + _HEAD_TAIL),
    ("skills", _HEAD_LEAD + r"((technical|core|key|professional)\s+)?(skills|competenc(y|ies)|technologies)" + _HEAD_TAIL),
    ("projects", _HEAD_LEAD + r"((key|notable|personal|academic|major|github|side|featured|selected)\s+)?projects?" + _HEAD_TAIL),
]

SECTION_PATTERNS = {
    "contact": r"(email|phone|linkedin|@)",
    "summary": r"(summary|objective|profile|professional summary)\b",
    "experience": r"(experience|employment|work history|internships?)\b",
    "education": r"(education|academic|academics)\b",
    "skills": r"(skills|technical skills|competencies|technologies)\b",
    "projects": r"(projects?)\b",
}

STOPWORDS = set("""
a an the and or of to in on for with as at by from is are was were be been
being this that these those it its your you i we our their his her they he
she will shall can could should would may might must not no nor so than then
""".split())

# Generic job-description boilerplate that shouldn't surface as a "missing keyword"
JD_FILLER_WORDS = set("""
need needs needed looking look strong required requirement requirements
experience experienced years year ability abilities skill skills work
working team teams role roles job jobs candidate candidates plus preferred
minimum environment including include includes excellent proven demonstrated
knowledge understanding familiarity responsible responsibilities
opportunity opportunities company companies join great good etc
""".split())

def extract_text(filepath, filename):
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        try:
            import pdfplumber
        except ImportError:
            raise ValueError(
                "PDF support isn't installed on this server. Run "
                "'pip install -r requirements.txt' (or just 'pip install pdfplumber') "
                "in the same Python environment used to start app.py, then restart the server."
            )
        text = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text.append(page.extract_text() or "")
        return "\n".join(text)
    elif ext == "docx":
        try:
            import docx
        except ImportError:
            raise ValueError(
                "DOCX support isn't installed on this server. Run "
                "'pip install -r requirements.txt' (or just 'pip install python-docx') "
                "in the same Python environment used to start app.py, then restart the server."
            )
        doc = docx.Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs)
    elif ext == "txt":
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: .{ext}")

def extract_contact_info(text):
    email = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
    # Try a labeled number ("Phone: ..."), then a separator-delimited number,
    # then a bare 10-13 digit run not touching other digits.
    phone = None
    labeled = re.search(
        r"(?:phone|mobile|cell|tel(?:ephone)?|contact(?:\s*(?:no|number))?|whatsapp)\s*[:.\-]?\s*"
        r"(\+?\d{1,3}[\s.-]?(?:\(\d{2,4}\)[\s.-]?)?\d[\d\s.-]{6,13}\d)",
        text,
        re.I,
    )
    if labeled:
        phone = labeled.group(1).strip()
    else:
        separated = re.search(
            r"(?<!\d)(\+?\d{1,3}[\s.-]?)?(\(\d{3}\)|\d{3})[\s.-]\d{3}[\s.-]?\d{4}(?!\d)",
            text,
        )
        if separated:
            phone = separated.group(0)
        else:
            bare = re.search(r"(?<!\d)(\+\d{1,3}[\s-]?)?\d{10,13}(?!\d)", text)
            if bare:
                phone = bare.group(0)
    linkedin = re.search(r"linkedin\.com/in/[\w-]+", text, re.I)
    return {
        "email": email.group(0) if email else None,
        "phone": phone,
        "linkedin": linkedin.group(0) if linkedin else None,
    }
def detect_sections(text):
    lower = text.lower()
    return {name: bool(re.search(pattern, lower)) for name, pattern in SECTION_PATTERNS.items()}

# FlashText Skill Extractor Setup
SKILL_SYNONYMS = {
    "React": ["react", "react.js", "reactjs"],
    "Python": ["python", "python3"],
    "Node.js": ["node.js", "node js", "nodejs", "node"],
    "Machine Learning": ["machine learning", "ml"],
    "Artificial Intelligence": ["artificial intelligence", "ai"],
    "C++": ["c++", "cpp"],
    "C#": ["c#", "csharp"],
    "Vue.js": ["vue", "vue.js", "vuejs"],
    "Next.js": ["next.js", "nextjs"],
    "Tailwind CSS": ["tailwind", "tailwindcss", "tailwind css"],
    "PostgreSQL": ["postgresql", "postgres"],
    "MongoDB": ["mongodb", "mongo"],
    "Amazon Web Services (AWS)": ["aws", "amazon web services"],
    "Google Cloud Platform (GCP)": ["gcp", "google cloud", "google cloud platform"],
    "Natural Language Processing": ["nlp", "natural language processing"],
    "Deep Learning": ["deep learning", "dl"],
    "Data Science": ["data science", "ds"]
}

keyword_processor = KeywordProcessor(case_sensitive=False)
for standard_name, synonyms in SKILL_SYNONYMS.items():
    keyword_processor.add_keywords_from_dict({standard_name: synonyms})
for category, skills in SKILL_DB.items():
    for skill in skills:
        found = False
        for standard, syns in SKILL_SYNONYMS.items():
            if skill.lower() in syns or skill.lower() == standard.lower():
                found = True
                break
        if not found:
            keyword_processor.add_keyword(skill, skill.title())
def extract_skills(text):
    return set(keyword_processor.extract_keywords(text))
def count_action_verb_bullets(text):
    # Reuses the bullet-detection heuristic from extract_bullet_lines()
    # below (resolved at call-time, so the ordering here is fine).
    bullets = extract_bullet_lines(text)
    total_bullets = 0
    strong_start = 0
    for clean in bullets:
        words = re.findall(r"[a-zA-Z']+", clean)
        if not words:
            continue
        total_bullets += 1
        if words[0].lower() in ACTION_VERBS:
            strong_start += 1
    return strong_start, total_bullets
def count_quantified_bullets(text):
    # Only counts numbers inside genuine bullet/achievement lines, not just
    # any line in the document (avoids counting dates in job-title lines).
    bullets = extract_bullet_lines(text)
    quantified = 0
    for clean in bullets:
        if re.search(r"\d", clean):
            quantified += 1
    return quantified
def find_weak_phrases(text):
    lower = text.lower()
    return [p for p in WEAK_PHRASES if p in lower]
def parse_job_description(jd_text):
    if not jd_text or not jd_text.strip():
        return None
    segments = {"Required": [], "Preferred": [], "Responsibilities": [], "Education": [], "Experience": []}
    lines = jd_text.splitlines()
    current_segment = "Required"
    for line in lines:
        lower = line.strip().lower()
        if not lower: continue
        if re.search(r'^(requirements|qualifications|must have|what you need)', lower):
            current_segment = "Required"
            continue
        elif re.search(r'^(nice to have|preferred|bonus|plus)', lower):
            current_segment = "Preferred"
            continue
        elif re.search(r'^(responsibilities|what you will do|duties)', lower):
            current_segment = "Responsibilities"
            continue
        if re.search(r'\b(bachelor|master|phd|degree)\b', lower):
            segments["Education"].append(line)
        if re.search(r'\d+[\+-]?\s*years?', lower):
            segments["Experience"].append(line)
        segments[current_segment].append(line)
    return {k: "\n".join(v) for k, v in segments.items()}
def extract_keywords_from_jd(jd_text, top_n=25):
    """Pull salient keywords out of a job description: known skills first,
    then top TF-IDF unigrams/bigrams as a fallback for anything not in the DB."""
    jd_skills = extract_skills(jd_text)
    try:
        vectorizer = TfidfVectorizer(
            stop_words="english", ngram_range=(1, 2), max_features=60
        )
        tfidf = vectorizer.fit_transform([jd_text])
        scores = tfidf.toarray()[0]
        terms = vectorizer.get_feature_names_out()
        ranked = sorted(zip(terms, scores), key=lambda x: x[1], reverse=True)
        extra_terms = [
            t for t, s in ranked
            if s > 0
            and t not in jd_skills
            # Skip TF-IDF terms that are themselves already-known skills (just
            # in a different casing/synonym form, e.g. "aws" vs the canonical
            # "Amazon Web Services (AWS)"), otherwise the same skill shows up
            # twice — once canonical, once raw — in matched/missing lists.
            and not extract_skills(t)
            and not any(w in JD_FILLER_WORDS for w in t.split())
        ][:top_n]
    except ValueError:
        extra_terms = []
    keywords = list(jd_skills) + extra_terms
    seen, ordered = set(), []
    for k in keywords:
        if k not in seen:
            seen.add(k)
            ordered.append(k)
    return ordered[:top_n]
def _keyword_in_text(keyword, text_lower):
    """Whole-word/whole-phrase containment check. Plain `keyword in text`
    substring checks cause false positives for short keywords — e.g. the JD
    keyword 'go' would "match" inside 'google' or 'going', and 'r' would
    match inside almost anything. This anchors on non-word/non +#. boundaries
    so only real, standalone occurrences count."""
    pattern = r"(?<![\w+#.])" + re.escape(keyword.lower()) + r"(?![\w+#])"
    return bool(re.search(pattern, text_lower))
def jd_match_score(resume_text, jd_text, resume_skills=None):
    # Model #1 - Sentence-BERT semantic resume <-> job description
    # similarity, with a TF-IDF fallback if the model isn't available.
    similarity = sbert_similarity(resume_text, jd_text)
    if similarity is None:
        vectorizer = TfidfVectorizer(stop_words="english")
        try:
            tfidf = vectorizer.fit_transform([resume_text, jd_text])
            similarity = cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0]
        except ValueError:
            similarity = 0.0
    jd_segments = parse_job_description(jd_text)
    if jd_segments and (jd_segments["Required"] or jd_segments["Preferred"]):
        req_skills = extract_keywords_from_jd(jd_segments["Required"])
        pref_skills = extract_keywords_from_jd(jd_segments["Preferred"])
    else:
        req_skills = extract_keywords_from_jd(jd_text)
        pref_skills = []
    resume_lower = resume_text.lower()
    if resume_skills is None:
        resume_skills = set()
    req_matched = [k for k in req_skills if k in resume_skills or _keyword_in_text(k, resume_lower)]
    req_missing = [k for k in req_skills if k not in req_matched]
    pref_matched = [k for k in pref_skills if k in resume_skills or _keyword_in_text(k, resume_lower)]
    pref_missing = [k for k in pref_skills if k not in pref_matched]
    all_matched = set(req_matched + pref_matched)
    keyword_density = {}
    for k in all_matched:
        pattern = r"(?<![\w+#.])" + re.escape(k.lower()) + r"(?![\w+#])"
        count = len(re.findall(pattern, resume_lower))
        if count == 0:
            count = resume_lower.count(k.lower())
        if count == 0 and k in resume_skills:
            count = 1
        keyword_density[k] = count
    return {
        "similarity": round(similarity * 100, 1),
        "required_matched": req_matched,
        "required_missing": req_missing,
        "preferred_matched": pref_matched,
        "preferred_missing": pref_missing,
        "keyword_density": keyword_density,
        "education_requirements": jd_segments["Education"][:300] if jd_segments and jd_segments["Education"] else "",
        "experience_requirements": jd_segments["Experience"][:300] if jd_segments and jd_segments["Experience"] else ""
    }

# Lighter-weight sibling of jd_match_score(): instead of a full job
# description, compares the resume against a curated skill list for a
# plainly-typed role title (e.g. "Data Analyst").
TARGET_ROLE_SKILLS = {
    "data analyst": ["SQL", "Excel", "Python", "R", "Tableau", "Power BI",
                      "Statistics", "Data Visualization", "A/B Testing", "Google Analytics"],
    "data scientist": ["Python", "R", "SQL", "Machine Learning", "Statistics", "Pandas",
                        "NumPy", "Scikit-learn", "TensorFlow", "PyTorch", "Data Visualization",
                        "Deep Learning"],
    "business analyst": ["SQL", "Excel", "Tableau", "Power BI", "Requirements Gathering",
                          "Stakeholder Management", "Business Process", "Agile", "JIRA"],
    "marketing analyst": ["Google Analytics", "Excel", "SQL", "A/B Testing", "SEO",
                           "Data Visualization", "Power BI"],
    "software engineer": ["Python", "Java", "C++", "JavaScript", "Git", "Data Structures",
                           "Algorithms", "SQL", "REST API", "Docker", "Agile"],
    "frontend developer": ["JavaScript", "TypeScript", "React", "HTML", "CSS", "Redux",
                            "Webpack", "Responsive Design", "Git"],
    "backend developer": ["Python", "Java", "Node.js", "SQL", "REST API", "Docker",
                           "Kubernetes", "PostgreSQL", "MongoDB", "Redis", "Microservices"],
    "full stack developer": ["JavaScript", "React", "Node.js", "HTML", "CSS", "SQL",
                              "REST API", "Git", "Docker"],
    "devops engineer": ["Docker", "Kubernetes", "Terraform", "Ansible", "Jenkins", "CI/CD",
                         "AWS", "Azure", "GCP", "Linux", "GitHub Actions", "Grafana"],
    "machine learning engineer": ["Python", "TensorFlow", "PyTorch", "Scikit-learn", "Pandas",
                                  "NumPy", "SQL", "Docker", "Kubernetes", "Machine Learning",
                                  "Deep Learning", "MLOps"],
    "product manager": ["Stakeholder Management", "Agile", "Scrum", "Roadmapping",
                         "Market Research", "Data Analysis", "JIRA", "Communication", "Leadership"],
    "project manager": ["Agile", "Scrum", "JIRA", "Stakeholder Management", "Risk Management",
                         "Budgeting", "Communication", "Leadership"],
    "ux designer": ["Figma", "Sketch", "Wireframing", "Prototyping", "User Research",
                     "Usability Testing", "Adobe XD"],
    "ui designer": ["Figma", "Sketch", "Adobe XD", "Prototyping", "Design Systems", "Typography"],
    "qa engineer": ["Selenium", "Manual Testing", "Automation Testing", "Test Cases", "JIRA",
                     "SQL", "API Testing", "Cypress"],
    "cloud engineer": ["AWS", "Azure", "GCP", "Terraform", "Kubernetes", "Docker", "Linux",
                        "Networking"],
    "cybersecurity analyst": ["Network Security", "SIEM", "Penetration Testing", "Firewalls",
                              "Vulnerability Assessment", "Python", "Linux"],
}

# Alternate phrasings that resolve to a canonical role key above.
TARGET_ROLE_ALIASES = {
    "data analytics": "data analyst",
    "data analysis": "data analyst",
    "business analytics": "business analyst",
    "ml engineer": "machine learning engineer",
    "ai engineer": "machine learning engineer",
    "swe": "software engineer",
    "sde": "software engineer",
    "front end developer": "frontend developer",
    "front-end developer": "frontend developer",
    "frontend engineer": "frontend developer",
    "back end developer": "backend developer",
    "back-end developer": "backend developer",
    "backend engineer": "backend developer",
    "fullstack developer": "full stack developer",
    "full-stack developer": "full stack developer",
    "devops": "devops engineer",
    "sre": "devops engineer",
    "site reliability engineer": "devops engineer",
    "pm": "product manager",
    "ui/ux designer": "ux designer",
    "ux/ui designer": "ux designer",
    "qa": "qa engineer",
    "qa tester": "qa engineer",
    "test engineer": "qa engineer",
    "sdet": "qa engineer",
    "cloud architect": "cloud engineer",
    "security analyst": "cybersecurity analyst",
    "infosec analyst": "cybersecurity analyst",
}

def match_target_role(role_name):
    """Resolve a free-typed role title to one of the canonical keys in
    TARGET_ROLE_SKILLS, or None if nothing close enough is found."""
    if not role_name or not role_name.strip():
        return None
    norm = re.sub(r"[^a-z0-9+#. ]", "", role_name.strip().lower())
    norm = re.sub(r"\s+", " ", norm).strip()
    norm = TARGET_ROLE_ALIASES.get(norm, norm)
    if norm in TARGET_ROLE_SKILLS:
        return norm
    keys = list(TARGET_ROLE_SKILLS.keys())
    close = difflib.get_close_matches(norm, keys, n=1, cutoff=0.8)
    if close:
        return close[0]

    # Substring match, e.g. "senior data analyst" -> "data analyst".
    for key in keys:
        if key in norm:
            return key
    return None
def target_role_skill_gap(resume_text, resume_skills, role_name):
    """Compare the resume against the expected skill set for a plainly typed
    target role (e.g. "Data Analyst") and report matched/missing skills."""
    canonical_role = match_target_role(role_name)
    if not canonical_role:
        return {
            "role_input": role_name,
            "recognized": False,
            "matched_role": None,
            "matched_skills": [],
            "missing_skills": [],
            "match_score": 0,
        }
    expected_skills = TARGET_ROLE_SKILLS[canonical_role]
    resume_lower = resume_text.lower()
    resume_skills_lower = {s.lower() for s in (resume_skills or set())}
    matched, missing = [], []
    for skill in expected_skills:
        if skill.lower() in resume_skills_lower or _keyword_in_text(skill, resume_lower):
            matched.append(skill)
        else:
            missing.append(skill)
    match_score = round(len(matched) / len(expected_skills) * 100) if expected_skills else 0
    return {
        "role_input": role_name,
        "recognized": True,
        "matched_role": canonical_role.title(),
        "matched_skills": matched,
        "missing_skills": missing,
        "match_score": match_score,
    }
def determine_experience_level(chunks):
    exp_text = chunks.get("experience", "")
    edu_text = chunks.get("education", "")
    if not exp_text.strip():
        return "Fresher"
    exp_words = len(exp_text.split())
    if exp_words < 60:
        return "Fresher"

    # Recent-grad window computed off today's date so it doesn't go stale.
    current_year = datetime.now().year
    recent_years = {str(y) for y in range(current_year - 2, current_year + 2)}
    years = [y for y in re.findall(r"(20\d{2})", edu_text) if y in recent_years]
    if years and exp_words < 150:
        return "Fresher"
    return "Experienced"
def score_projects(project_text):
    if not project_text.strip():
        return {"score": 0, "metrics": {}, "suggestions": ["No projects section found."]}
    skills_in_projects = extract_skills(project_text)
    word_count = len(project_text.split())
    strong_bullets, total_bullets = count_action_verb_bullets(project_text)
    quantified = count_quantified_bullets(project_text)
    tech_score = min(len(skills_in_projects) * 15, 100)
    impact_score = min(quantified * 20, 100)
    complexity_score = min(word_count / 1.5, 100)
    action_score = (strong_bullets / total_bullets * 100) if total_bullets else 0
    overall = round(tech_score * 0.3 + impact_score * 0.3 + complexity_score * 0.2 + action_score * 0.2)
    suggestions = []
    if tech_score < 50:
        suggestions.append("Explicitly mention the technologies, frameworks, and libraries used in your projects.")
    if impact_score < 50:
        suggestions.append("Quantify the impact of your projects (e.g., 'served 500 users', 'improved speed by 20%').")
    if action_score < 50:
        suggestions.append("Start project descriptions with strong action verbs (e.g., 'Architected', 'Developed').")
    return {
        "score": overall,
        "metrics": {
            "tech_stack_score": round(tech_score),
            "impact_score": round(impact_score),
            "complexity_score": round(complexity_score),
            "action_verbs_score": round(action_score)
        },
        "suggestions": suggestions if suggestions else ["Great project descriptions with solid technical depth and measurable impact."]
    }
def score_resume(text, jd_text=None, ats_risk=None, target_role=None):
    word_count = len(re.findall(r"\w+", text))
    sections = detect_sections(text)
    contact = extract_contact_info(text)
    skills_found = extract_skills(text)
    strong_bullets, total_bullets = count_action_verb_bullets(text)
    quantified = count_quantified_bullets(text)
    weak_phrases = find_weak_phrases(text)

    # Structure score
    structure_score = 0
    structure_checks = []
    if contact["email"]:
        structure_score += 20
        structure_checks.append(("Email found", True))
    else:
        structure_checks.append(("Email found", False))
    if contact["phone"]:
        structure_score += 10
        structure_checks.append(("Phone number found", True))
    else:
        structure_checks.append(("Phone number found", False))
    for key in ["experience", "education", "skills"]:
        if sections[key]:
            structure_score += 15
        structure_checks.append((f"'{key.capitalize()}' section present", sections[key]))
    if sections["summary"]:
        structure_score += 10
    structure_checks.append(("Summary/objective present", sections["summary"]))
    if 350 <= word_count <= 900:
        structure_score += 15
        structure_checks.append(("Resume length appropriate (350-900 words)", True))
    else:
        structure_checks.append(("Resume length appropriate (350-900 words)", False))
    ats_penalty = 0
    if ats_risk:
        if ats_risk["risk_level"] == "High":
            ats_penalty = 20
        elif ats_risk["risk_level"] == "Medium":
            ats_penalty = 10
    structure_score = min(structure_score, 100)
    structure_score = max(0, structure_score - ats_penalty)
    if ats_penalty > 0:
        structure_checks.append((f"ATS formatting penalty (-{ats_penalty} pts)", False))
    else:
        structure_checks.append(("ATS formatting clean", True))

    # Content quality score
    content_score = 0
    bullet_ratio = (strong_bullets / total_bullets) if total_bullets else 0
    content_score += round(bullet_ratio * 40)
    quant_ratio = min(quantified / max(total_bullets, 1), 1)
    content_score += round(quant_ratio * 35)
    penalty = min(len(weak_phrases) * 5, 25)
    content_score += (25 - penalty)
    content_score = max(0, min(content_score, 100))

    # Skills coverage score
    skills_score = min(len(skills_found) * 6, 100)

    # Weights differ by experience level
    chunks = split_sections(text)
    exp_level = determine_experience_level(chunks)
    project_quality = score_projects(chunks.get("projects", ""))
    if exp_level == "Fresher":
        w_struct, w_content, w_skills, w_jd = 0.25, 0.30, 0.25, 0.20
        w_struct_nj, w_content_nj, w_skills_nj = 0.30, 0.40, 0.30
    else:
        w_struct, w_content, w_skills, w_jd = 0.20, 0.40, 0.25, 0.15
        w_struct_nj, w_content_nj, w_skills_nj = 0.20, 0.50, 0.30
    result = {
        "project_quality": project_quality,
        "word_count": word_count,
        "contact": contact,
        "sections": sections,
        "structure_checks": structure_checks,
        "structure_score": structure_score,
        "content_score": content_score,
        "skills_score": skills_score,
        "skills_found": sorted(skills_found),
        "strong_action_bullets": strong_bullets,
        "total_bullets_detected": total_bullets,
        "quantified_bullets": quantified,
        "weak_phrases": weak_phrases,
    }

    # Optional: JD match
    if jd_text and jd_text.strip():
        match = jd_match_score(text, jd_text, resume_skills=skills_found)
        result["jd_match"] = match
        overall = round(
            structure_score * w_struct
            + content_score * w_content
            + skills_score * w_skills
            + match["similarity"] * w_jd
        )
    else:
        result["jd_match"] = None
        overall = round(structure_score * w_struct_nj + content_score * w_content_nj + skills_score * w_skills_nj)
    result["overall_score"] = max(0, min(overall, 100))

    # Optional: target-role skill gap
    if target_role and target_role.strip():
        result["target_role_match"] = target_role_skill_gap(text, skills_found, target_role)
    else:
        result["target_role_match"] = None

    result["suggestions"] = build_suggestions(result)

    result["section_scores"] = section_wise_scores(text)
    result["completeness"] = completeness_score(
        text, sections, contact, skills_found, quantified, total_bullets
    )

    # Personal info / bias-risk check — advisory only, doesn't affect any score.
    result["personal_info_risk"] = detect_personal_info(
        text, has_photo=bool((ats_risk or {}).get("has_photo"))
    )

    # spaCy: structured extraction of education/experience/certifications,
    # plus an NER-based skills cross-check. None if spaCy isn't installed.
    entities = extract_entities_spacy(
        text,
        skill_vocab=ALL_SKILLS,
        experience_text=chunks.get("experience"),
        education_text=chunks.get("education"),
    )
    result["entities"] = entities
    if entities and entities.get("ner_skills"):
        merged = sorted(set(result["skills_found"]) | set(entities["ner_skills"]))
        result["skills_found"] = merged

    result["duplicate_content"] = detect_duplicate_content(text, chunks=chunks)

    # Resume text highlighting: prefer JD-matched/missing keywords when a JD
    # was supplied, then target-role matched/missing skills, and fall back
    # to the general skill database matches with no "missing" list when
    # neither was given.
    highlight_keywords = set(result["skills_found"])
    highlight_missing = []
    if result.get("jd_match"):
        jm = result["jd_match"]
        highlight_keywords |= set(jm.get("required_matched", [])) | set(jm.get("preferred_matched", []))
        highlight_missing = list(jm.get("required_missing", [])) + list(jm.get("preferred_missing", []))
    elif result.get("target_role_match") and result["target_role_match"].get("recognized"):
        trm = result["target_role_match"]
        highlight_keywords |= set(trm.get("matched_skills", []))
        highlight_missing = list(trm.get("missing_skills", []))
    result["text_highlights"] = build_resume_highlights(
        text, keywords_to_highlight=highlight_keywords, missing_keywords=highlight_missing
    )
    return result
def build_suggestions(r):
    tips = []
    if not r["contact"]["email"]:
        tips.append("Add a professional email address near the top of your resume.")
    if not r["contact"]["phone"]:
        tips.append("Include a phone number so recruiters can reach you quickly.")
    if not r["sections"]["summary"]:
        tips.append("Add a 2-3 line summary at the top tailored to the role you want.")
    if not r["sections"]["skills"]:
        tips.append("Add a dedicated 'Skills' section so ATS systems can parse your keywords.")
    if r["word_count"] < 350:
        tips.append("Your resume looks short — add more detail on impact and responsibilities.")
    elif r["word_count"] > 900:
        tips.append("Your resume is quite long — trim it down to 1-2 pages of the most relevant content.")
    if r["total_bullets_detected"] > 0:
        ratio = r["strong_action_bullets"] / r["total_bullets_detected"]
        if ratio < 0.5:
            tips.append(
                "Start more bullet points with strong action verbs "
                "(e.g. 'Led', 'Built', 'Reduced') instead of passive phrasing."
            )
    if r["quantified_bullets"] < max(r["total_bullets_detected"] // 2, 1):
        tips.append(
            "Quantify your achievements with numbers, percentages, or dollar amounts "
            "(e.g. 'Reduced load time by 40%')."
        )
    if r["weak_phrases"]:
        tips.append(
            "Replace overused phrases like "
            + ", ".join(f"'{p}'" for p in r["weak_phrases"][:3])
            + " with specific, evidence-backed statements."
        )
    if len(r["skills_found"]) < 6:
        tips.append("List more relevant technical and soft skills explicitly — ATS software scans for exact keyword matches.")
    if r["jd_match"]:
        missing = (r["jd_match"].get("required_missing", []) + r["jd_match"].get("preferred_missing", []))[:8]
        if missing:
            tips.append(
                "The job description mentions these keywords that don't appear in your resume: "
                + ", ".join(missing)
                + ". Add the ones that genuinely apply to your experience."
            )
        if r["jd_match"]["similarity"] < 50:
            tips.append("Your resume's overall similarity to this job description is low — mirror its terminology where truthful.")
    if r.get("target_role_match"):
        trm = r["target_role_match"]
        if trm["recognized"] and trm["missing_skills"]:
            tips.append(
                f"For a {trm['matched_role']} role, your resume is missing: "
                + ", ".join(trm["missing_skills"][:8])
                + ". Add the ones that genuinely apply to your experience."
            )
        elif not trm["recognized"] and trm["role_input"]:
            tips.append(
                f"'{trm['role_input']}' wasn't recognized as a target role — try a more common title "
                "(e.g. 'Data Analyst', 'Software Engineer', 'Product Manager')."
            )
    if not tips:
        tips.append("Great work — your resume covers the fundamentals well. Consider a final proofread and a tailored summary per application.")
    return tips

# Section-wise scoring

def split_sections(text):
    """Best-effort split of the resume into named chunks, based on section
    headers that appear alone on their own line. Anything before the first
    recognised header is treated as the 'header' zone (name/contact/summary
    area). This is heuristic — resumes with unusual formatting may not split
    perfectly, but it's good enough to give section-level signal."""
    lines = text.splitlines()
    sections = {}
    current = "header"
    buffer = []
    def flush():
        sections.setdefault(current, [])
        sections[current].extend(buffer)
    for line in lines:
        stripped = line.strip()
        matched = None
        for name, pattern in SECTION_HEADER_PATTERNS:
            if re.match(pattern, stripped, re.I):
                matched = name
                break
        if matched:
            flush()
            buffer = []
            current = matched
        else:
            buffer.append(line)
    flush()
    return {k: "\n".join(v) for k, v in sections.items()}
def section_wise_scores(text):
    """Score each major resume section independently (0-100), so a person
    can see exactly which part of the resume is weakest rather than just an
    overall number."""
    chunks = split_sections(text)
    scores = {}

    # Contact info can sit anywhere near the top, so pull from the whole doc.
    contact = extract_contact_info(text)
    c_score = 0
    if contact["email"]:
        c_score += 50
    if contact["phone"]:
        c_score += 30
    if contact["linkedin"]:
        c_score += 20
    scores["contact"] = {"label": "Contact info", "score": min(c_score, 100), "present": c_score > 0}

    # Summary
    summary_text = chunks.get("summary", "").strip()
    if summary_text:
        wc = len(summary_text.split())
        s_score = 100 if 20 <= wc <= 80 else (65 if wc > 0 else 0)
        scores["summary"] = {"label": "Summary / objective", "score": s_score, "present": True}
    else:
        scores["summary"] = {"label": "Summary / objective", "score": 0, "present": False}

    # Experience
    exp_text = chunks.get("experience", "")
    if exp_text.strip():
        strong, total = count_action_verb_bullets(exp_text)
        quant = count_quantified_bullets(exp_text)
        ratio_action = (strong / total) if total else 0
        ratio_quant = min(quant / max(total, 1), 1)
        e_score = round(ratio_action * 60 + ratio_quant * 40)
        scores["experience"] = {
            "label": "Experience", "score": e_score, "present": True, "bullets_detected": total,
        }
    else:
        scores["experience"] = {"label": "Experience", "score": 0, "present": False, "bullets_detected": 0}

    # Education
    edu_text = chunks.get("education", "")
    if edu_text.strip():
        has_year = bool(re.search(r"(19|20)\d{2}", edu_text))
        has_degree = bool(re.search(r"(bachelor|master|b\.?s\.?|m\.?s\.?|ph\.?d|associate|diploma|degree)", edu_text, re.I))
        ed_score = 40 + (30 if has_year else 0) + (30 if has_degree else 0)
        scores["education"] = {"label": "Education", "score": ed_score, "present": True}
    else:
        scores["education"] = {"label": "Education", "score": 0, "present": False}

    # Skills
    skills_text = chunks.get("skills", "")
    if skills_text.strip():
        found = extract_skills(skills_text)
        sk_score = min(len(found) * 10, 100)
        scores["skills"] = {"label": "Skills", "score": sk_score, "present": True, "count": len(found)}
    else:
        found = extract_skills(text)
        sk_score = min(len(found) * 5, 60) if found else 0
        scores["skills"] = {"label": "Skills", "score": sk_score, "present": bool(found), "count": len(found)}

    # Projects (optional section)
    proj_text = chunks.get("projects", "")
    if proj_text.strip():
        strong, total = count_action_verb_bullets(proj_text)
        p_score = round((strong / total) * 100) if total else 55
        scores["projects"] = {"label": "Projects", "score": p_score, "present": True}
    else:
        scores["projects"] = {"label": "Projects", "score": None, "present": False}
    return scores
def completeness_score(text, sections, contact, skills_found, quantified, total_bullets):
    has_github = bool(re.search(r"github\.com/[\w-]+", text, re.I))
    has_certifications = bool(re.search(r"^\s*(certifications?|licenses?)\s*:?\s*$", text, re.I | re.M))
    has_achievements = bool(re.search(r"^\s*(achievements?|awards?|honors?|recognitions?)\s*:?\s*$", text, re.I | re.M))
    checklist = [
        ("Summary", sections["summary"]),
        ("Skills", sections["skills"]),
        ("Experience", sections["experience"]),
        ("Projects", sections["projects"]),
        ("Education", sections["education"]),
        ("Certifications", has_certifications),
        ("Achievements", has_achievements),
        ("Contact Information", bool(contact["email"] or contact["phone"])),
        ("LinkedIn", bool(contact["linkedin"])),
        ("GitHub", has_github),
    ]
    passed = sum(1 for _, ok in checklist if ok)
    score = round(passed / len(checklist) * 100)
    missing = [label for label, ok in checklist if not ok]
    return {"score": score, "checklist": checklist, "missing": missing}

# ATS risk analysis

UNUSUAL_BULLET_CHARS = re.compile(r"[➤◆■●▪✦❖✔➔]")

def analyze_ats_risk(filepath, filename, text):
    """Heuristic scan for formatting choices that commonly trip up ATS
    (Applicant Tracking System) parsers: tables, embedded images, multi-column
    layouts, exotic bullet glyphs, and missing standard section headers.
    This can't see the actual visual layout/fonts, only what the parser
    itself could extract — which is exactly the same limitation a real ATS
    has, so it's a reasonable proxy."""
    ext = filename.rsplit(".", 1)[-1].lower()
    issues = []
    has_tables = False
    has_images = False
    try:
        if ext == "pdf":
            import pdfplumber
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    try:
                        if page.find_tables():
                            has_tables = True
                    except Exception:
                        pass
                    if page.images:
                        has_images = True
        elif ext == "docx":
            import docx
            doc = docx.Document(filepath)
            if doc.tables:
                has_tables = True
            if doc.inline_shapes:
                has_images = True
    except Exception:
        pass
    if has_tables:
        issues.append({
            "issue": "Tables detected",
            "severity": "high",
            "tip": "Avoid tables for layout — many ATS parsers read table cells out of "
                   "order or skip them entirely. Use plain single-column text with clear headers instead.",
        })
    if has_images:
        issues.append({
            "issue": "Images or graphics detected",
            "severity": "medium",
            "tip": "Avoid logos, icons, headshots, or charts — ATS software cannot read "
                   "image content, and it can break parsing of nearby text.",
        })

    # Multi-column heuristic: lines with large internal gaps often indicate
    # side-by-side columns that a linear text extractor reads out of order.
    gap_lines = sum(1 for l in text.splitlines() if re.search(r"\S {4,}\S", l))
    if gap_lines > 5:
        issues.append({
            "issue": "Possible multi-column layout",
            "severity": "medium",
            "tip": "Several lines show large internal gaps, which often means a "
                   "multi-column layout. Many ATS systems read columns left-to-right "
                   "across the whole page and scramble the content. A single column is safest.",
        })
    if UNUSUAL_BULLET_CHARS.search(text):
        issues.append({
            "issue": "Non-standard bullet or icon symbols",
            "severity": "low",
            "tip": "Stick to simple bullets (- or •). Exotic symbols and icon fonts can "
                   "render as garbled characters or boxes when an ATS extracts the text.",
        })
    matched_sections = sum(
        1 for name in ["experience", "education", "skills"]
        if re.search(SECTION_PATTERNS[name], text.lower())
    )
    if matched_sections < 3:
        issues.append({
            "issue": "Non-standard or missing section headers",
            "severity": "high",
            "tip": "Use conventional headers such as 'Experience', 'Education', and "
                   "'Skills' on their own line. Creative header names (e.g. 'My Journey') "
                   "often aren't recognised by ATS section parsers.",
        })
    if not text.strip():
        issues.append({
            "issue": "No extractable text",
            "severity": "high",
            "tip": "This file appears to contain no machine-readable text (it may be a "
                   "scanned image). ATS systems cannot read scanned resumes at all — export as a text-based PDF or DOCX.",
        })
    high = sum(1 for i in issues if i["severity"] == "high")
    medium = sum(1 for i in issues if i["severity"] == "medium")
    if high:
        level = "High"
    elif medium:
        level = "Medium"
    else:
        level = "Low"

    # XGBoost: a learned 0-100 ATS-parseability score alongside the
    # rule-based risk_level/issues above.
    contact = extract_contact_info(text)
    sections = detect_sections(text)
    skills_found = extract_skills(text)
    strong_bullets, total_bullets = count_action_verb_bullets(text)
    quantified = count_quantified_bullets(text)
    num_core_sections = sum(1 for k in ["experience", "education", "skills", "summary"] if sections.get(k))
    features = build_ats_feature_vector(
        word_count=len(re.findall(r"\w+", text)),
        has_email=bool(contact["email"]),
        has_phone=bool(contact["phone"]),
        has_linkedin=bool(contact["linkedin"]),
        num_core_sections=num_core_sections,
        has_tables=has_tables,
        has_images=has_images,
        gap_line_ratio=(gap_lines / max(len(text.splitlines()), 1)),
        has_unusual_bullets=bool(UNUSUAL_BULLET_CHARS.search(text)),
        matched_standard_headers=matched_sections,
        skills_count=len(skills_found),
        bullet_count=total_bullets,
        quantified_bullet_ratio=(quantified / max(total_bullets, 1)),
    )
    ml_score = predict_ats_score(features)
    return {
        "risk_level": level,
        "issues": issues,
        "ml_ats_score": ml_score,
        # Surfaced separately (not just folded into `issues`) so the personal-info
        # privacy check below can flag "photo present" as its own item without
        # having to re-open and re-parse the PDF/DOCX file a second time.
        "has_photo": has_images,
    }

# Personal info / bias-risk check
#
# Flags personal or demographic details that are common on resume templates
# in some regions (a photo, date of birth, marital status, etc.) but that
# many companies hiring in the US/UK/Canada/EU specifically don't want to
# see: several of these are protected characteristics, so some employers'
# ATS platforms and hiring policies are configured to flag or auto-reject
# applications that include them rather than take on discrimination risk.
# Purely advisory — this never changes any existing score.

PERSONAL_INFO_CHECKS = [
    ("dob_age", "Date of birth / age", "high",
     r"\b(date\s+of\s+birth|d\.?\s?o\.?\s?b\.?)\s*[:.\-]|\bborn\s+(on|in)\s+\d|"
     r"\b\d{1,2}\s*years?\s*old\b|\bage\s*[:\-]\s*\d{1,2}\b",
     "Age and date of birth are protected characteristics in most hiring regions "
     "(US, UK, EU, Canada) — including one can expose an employer to age-"
     "discrimination risk, so some will discard a resume that lists it rather than "
     "take that on. Your work history already signals seniority; leave the exact "
     "age or DOB off."),
    ("marital_status", "Marital / family status", "medium",
     r"\bmarital\s+status\b|\b(unmarried|divorced|widow(?:ed)?)\b",
     "Marital or family status has no bearing on job performance and is a "
     "protected category in most Western hiring markets — it's safest left off a "
     "resume aimed at those employers."),
    ("nationality", "Nationality / citizenship", "low",
     r"\b(nationality|citizenship)\s*[:.\-]",
     "Only include nationality or work-authorization status if the job posting "
     "specifically asks about it — otherwise it's personal data an ATS has no "
     "field for and a recruiter doesn't need."),
    ("gender", "Gender", "medium",
     r"\b(gender|sex)\s*[:.\-]\s*(male|female|m|f)\b",
     "Gender is a protected characteristic and isn't relevant to qualifications — "
     "some employers flag its presence on a resume as a compliance concern during screening."),
    ("religion", "Religion", "medium",
     r"\breligion\s*[:.\-]",
     "Religious affiliation is a protected characteristic in most hiring markets "
     "and has no bearing on job performance — leave it off."),
]

def detect_personal_info(text, has_photo=False):
    """Scans for personal/demographic fields that show up on resume
    templates in some regions but that many companies' ATS platforms and
    hiring policies specifically don't want to see. Returns a flat list of
    flags plus an overall risk read, in the same shape as analyze_ats_risk()
    so the UI can reuse the same components to render it."""
    flags = []
    if has_photo:
        flags.append({
            "item": "Photo / headshot",
            "severity": "medium",
            "tip": "Many US/UK/Canadian employers explicitly ask candidates not to "
                   "include a photo — it can introduce unconscious bias during human "
                   "screening, and most ATS parsers discard the image anyway. Safe to "
                   "remove unless the region or industry you're applying to expects one.",
        })
    for _, label, severity, pattern, tip in PERSONAL_INFO_CHECKS:
        if re.search(pattern, text, re.I):
            flags.append({"item": label, "severity": severity, "tip": tip})

    high = sum(1 for f in flags if f["severity"] == "high")
    medium = sum(1 for f in flags if f["severity"] == "medium")
    if high:
        level = "High"
    elif medium:
        level = "Medium"
    elif flags:
        level = "Low"
    else:
        level = "Clean"
    return {
        "flags": flags,
        "risk_level": level,
        "has_issues": len(flags) > 0,
    }

# Bullet point rewriting

_BULLET_MARKER_RE = re.compile(r"^[•\-\*\u2022➤◆■●▪]")

def _merge_wrapped_lines(raw_lines):
    """pdfplumber/docx extraction returns one physical line per visually
    wrapped line, so a single bullet that wraps onto a second line in the
    original resume (very common for anything ending in a quantified
    result, e.g. "...delivered predictions with" / "95% accuracy.") comes
    back here as two separate lines with no bullet marker on the second
    one. Left unmerged, the first (incomplete) line gets treated as its
    own weak bullet — truncated mid-sentence — and the orphaned
    continuation line is silently dropped, which also makes a genuinely
    quantified bullet get flagged as having no measurable outcome.

    This reassembles a continuation line onto the previous line only when
    both signals agree: the previous line doesn't already end in
    sentence-ending punctuation, AND the current line starts lowercase or
    with a digit (e.g. "95%", "$10K") — the two hallmarks of a wrapped
    trailing clause rather than a genuinely new bullet/header."""
    merged = []
    for line in raw_lines:
        is_marked = bool(_BULLET_MARKER_RE.match(line))
        is_shouty_header = line.isupper() and len(line.split()) <= 6
        prev_incomplete = merged and not merged[-1].rstrip().endswith((".", ":", "!", "?"))
        looks_like_continuation = bool(re.match(r"^[a-z]", line)) or bool(re.match(r"^\d", line))
        if (
            merged
            and not is_marked
            and not is_shouty_header
            and "@" not in line and "|" not in line
            and prev_incomplete
            and looks_like_continuation
        ):
            merged[-1] = merged[-1].rstrip() + " " + line.strip()
        else:
            merged.append(line)
    return merged

def extract_bullet_lines(text):
    """Pulls out lines that actually look like resume bullet points — either
    explicitly marked (-, *, •, ...) or sentence-like lines with enough lowercase
    connective words to distinguish them from short Title Case headers like a
    job title or company name line."""
    raw_lines = [l.strip() for l in text.splitlines() if l.strip()]
    lines = _merge_wrapped_lines(raw_lines)
    bullets = []
    for l in lines:
        is_marked = bool(_BULLET_MARKER_RE.match(l))
        clean = re.sub(r"^[•\-\*\u2022➤◆■●▪]\s*", "", l).strip()
        if "@" in clean or "|" in clean:
            continue
        words = clean.split()
        if is_marked:
            if clean:
                bullets.append(clean)
            continue
        if 5 <= len(words) <= 30:
            lowercase_after_first = sum(1 for w in words[1:] if w[:1].islower())
            if lowercase_after_first >= 2:
                bullets.append(clean)
    return bullets
def is_weak_bullet(bullet):
    words = re.findall(r"[a-zA-Z']+", bullet)
    if not words:
        return True
    starts_weak = words[0].lower() not in ACTION_VERBS
    has_number = bool(re.search(r"\d", bullet))
    return starts_weak or not has_number
def suggest_verb(bullet):
    lower = bullet.lower()
    for pattern, verb in VERB_HINTS:
        if re.search(pattern, lower):
            return verb
    return DEFAULT_VERB
def rule_based_bullet_rewrite(bullet, is_project=False):
    """Rewrites a single weak bullet: adds a strong leading action verb if
    missing, and a quantified-result prompt if missing. For bullets sourced
    from the Projects section, also nudges toward naming a tech stack, since
    that's a signal recruiters specifically look for there but not in
    general Experience bullets."""
    words = re.findall(r"[a-zA-Z']+", bullet)
    starts_weak = not words or words[0].lower() not in ACTION_VERBS
    has_number = bool(re.search(r"\d", bullet))
    mentions_tech = None
    if is_project:
        # Reuses the shared skill vocabulary rather than a separate keyword list.
        mentions_tech = bool(extract_skills(bullet)) or any(
            kw in bullet.lower() for kw in ["using", "developed with", "built with", "technologies", "stack"]
        )
    core = WEAK_OPENERS_RE.sub("", bullet).strip()
    notes = []
    if starts_weak:
        verb = suggest_verb(bullet)
        # Only lowercase when we're prepending a new leading verb.
        if core:
            core = core[0].lower() + core[1:]
        core = f"{verb} {core}".strip()
        notes.append("Now leads with a strong action verb instead of a passive/weak phrase.")
    if is_project and not mentions_tech:
        core = core.rstrip(". ") + ", built with [add your tech stack/tools here]"
        notes.append("Explicitly name the technologies, frameworks, or tools used in this project.")
    if not has_number:
        core = core.rstrip(". ") + " — add a measurable result (e.g. 'by 25%', 'saving $10K/year', 'for 200+ users')."
        notes.append("No quantified outcome detected — add a specific number, percentage, or dollar amount.")
    if not notes:
        notes.append("Already reasonably strong — consider tightening the wording further.")
    if is_project and not core.endswith("."):
        core = core.rstrip(", ") + "."
    return {
        "original": bullet,
        "suggested": core,
        "notes": notes,
        "section": "Project" if is_project else "Experience",
    }
def get_bullet_rewrites(text, max_bullets=6):
    """Returns rule-based rewrite suggestions for the weakest bullet points
    found across the resume's Experience and Projects sections (weak opener
    and/or no quantified result). Each suggestion is tagged with the section
    it came from, and Project bullets get an extra tech-stack check that
    doesn't apply to Experience bullets."""
    chunks = split_sections(text)
    exp_bullets = extract_bullet_lines(chunks.get("experience", ""))
    proj_bullets = extract_bullet_lines(chunks.get("projects", ""))
    if not exp_bullets and not proj_bullets:
        # Fallback if section splitting didn't find headers: treat the
        # whole resume as one undifferentiated source (no tech-stack check).
        exp_bullets = extract_bullet_lines(text)

    weak = [(b, False) for b in exp_bullets if is_weak_bullet(b)]
    weak += [(b, True) for b in proj_bullets if is_weak_bullet(b)]
    weak = weak[:max_bullets]
    if not weak:
        return []
    return [rule_based_bullet_rewrite(b, is_project=is_proj) for b, is_proj in weak]

# Resume text highlighting
#
# Renders the resume's own text back with weak sentences/bullets and
# recognized keywords marked inline, so a person can see exactly *where* on
# the page a problem sentence or a matched skill lives, instead of only
# seeing it summarized as a number or a chip elsewhere in the report.
# Missing keywords obviously can't be marked inside text that doesn't
# contain them, so those are returned as a separate list for the UI to
# display alongside the highlighted resume, in their own distinct color.

def _bullet_is_weak(clean_line):
    """Same weak-writing signal used elsewhere in this file (no strong
    opening verb, no quantified result, or a known weak/filler phrase), but
    as a single yes/no check against one line instead of an aggregate ratio."""
    if not clean_line.strip():
        return False
    lower = clean_line.lower()
    if any(p in lower for p in WEAK_PHRASES):
        return True
    if WEAK_OPENERS_RE.match(clean_line.strip()):
        return True
    words = re.findall(r"[a-zA-Z']+", clean_line)
    first_is_verb = bool(words) and words[0].lower() in ACTION_VERBS
    has_number = bool(re.search(r"\d", clean_line))
    return not first_is_verb and not has_number

def _find_weak_line_spans(text):
    """Walks the resume line by line, tracking absolute character offsets,
    and returns (start, end) spans for every bullet/sentence-like line that
    reads as weak — so the exact original wording can be highlighted in place."""
    spans = []
    offset = 0
    for line in text.splitlines(keepends=True):
        stripped = line.rstrip("\n").rstrip("\r")
        content = stripped.strip()
        if not content:
            offset += len(line)
            continue
        is_marked_bullet = bool(_BULLET_MARKER_RE.match(content))
        word_count = len(content.split())
        looks_sentence_like = word_count >= 6 and re.search(r"[a-z]{3,}", content.lower())
        if is_marked_bullet or looks_sentence_like:
            clean = re.sub(r"^[•\-\*\u2022➤◆■●▪]\s*", "", content)
            if _bullet_is_weak(clean):
                # Locate `clean` (the bullet with its marker/whitespace
                # stripped off) within this exact physical line, so the
                # highlighted span lines up with the original text.
                pos = line.find(clean)
                abs_start = offset + pos if pos != -1 else offset
                abs_end = abs_start + len(clean)
                spans.append((abs_start, abs_end))
        offset += len(line)
    return spans

def _find_weak_phrase_char_spans(text):
    """Character-level spans for known weak/filler phrases anywhere in the
    document, not just inside a qualifying bullet/sentence-length line — so
    a buzzword sitting in a short summary line (e.g. "Hardworking team
    player") still gets highlighted even though the line itself is too
    short to be caught by _find_weak_line_spans()."""
    spans = []
    lower = text.lower()
    for phrase in WEAK_PHRASES:
        for m in re.finditer(re.escape(phrase), lower):
            spans.append((m.start(), m.end()))
    return spans

def _find_keyword_spans(text, keywords):
    """Case-insensitive, whole-word/phrase occurrences of each keyword in
    the resume text, for inline highlighting."""
    spans = []
    lower_text = text.lower()
    for kw in keywords:
        if not kw or len(kw) < 2:
            continue
        pattern = r"(?<![\w+#.])" + re.escape(kw.lower()) + r"(?![\w+#])"
        for m in re.finditer(pattern, lower_text):
            spans.append((m.start(), m.end()))
    return spans

def build_resume_highlights(text, keywords_to_highlight=None, missing_keywords=None):
    """Returns {segments, missing_keywords, ...} where `segments` is an
    ordered list of {text, type} chunks that concatenate back into the
    exact original resume text, each tagged "weak", "keyword", or "plain"
    so the frontend can render the resume itself with inline highlighting
    rather than only aggregate scores/chips."""
    keywords_to_highlight = list(keywords_to_highlight or [])
    missing_keywords = list(missing_keywords or [])

    weak_spans = [(s, e, "weak") for s, e in _find_weak_line_spans(text)]
    weak_spans += [(s, e, "weak") for s, e in _find_weak_phrase_char_spans(text)]
    keyword_spans = [(s, e, "keyword") for s, e in _find_keyword_spans(text, keywords_to_highlight)]

    # Weak spans win over keyword spans on overlap — a matched keyword
    # sitting inside a flagged weak sentence still reads as part of that
    # sentence, so the whole thing is shown as weak rather than split up.
    tagged = sorted(
        weak_spans + keyword_spans,
        key=lambda t: (t[0], 0 if t[2] == "weak" else 1, -(t[1] - t[0])),
    )
    merged = []
    cursor = 0
    for start, end, kind in tagged:
        if start < cursor:
            continue
        merged.append((start, end, kind))
        cursor = end

    segments = []
    pos = 0
    for start, end, kind in merged:
        if start > pos:
            segments.append({"text": text[pos:start], "type": "plain"})
        segments.append({"text": text[start:end], "type": kind})
        pos = end
    if pos < len(text):
        segments.append({"text": text[pos:], "type": "plain"})

    return {
        "segments": segments,
        "missing_keywords": sorted(set(missing_keywords))[:24],
        "weak_line_count": len(weak_spans),
        "keyword_match_count": len(keyword_spans),
    }

# Duplicate Content Detection
#
# Flags redundancy that a rule-based scorer above wouldn't otherwise catch:
# skills listed twice in the Skills section (usually a copy-paste slip),
# skills mentioned an unusually high number of times across the whole
# document (possible keyword stuffing), sentences repeated verbatim, and
# bullet points that are duplicated or near-duplicated across different
# jobs/projects (a common tell of a resume stitched together in a hurry).
# This is purely a reporting layer — it doesn't change structure_score,
# content_score, skills_score, or any other existing scoring path.

SKILL_SEPARATOR_RE = re.compile(r"[,\n\r|;•\u2022\u25CF\u25AA]+")
SENTENCE_SPLIT_RE = re.compile(r"(?<=[.!?])\s+")

def find_duplicate_skills(chunks):
    """Flags skills listed more than once, verbatim, inside the resume's own
    Skills section. Splits on commas/newlines/bullets/pipes (not hyphens or
    slashes, so compound names like 'Node.js' or 'CI/CD' survive intact),
    and strips a leading sub-header word (e.g. 'Languages:') off the first
    token of each line so that doesn't get counted as a skill itself."""
    skills_text = (chunks.get("skills") or "").strip()
    if not skills_text:
        return []
    raw_tokens = [t.strip() for t in SKILL_SEPARATOR_RE.split(skills_text) if t.strip()]
    tokens = []
    for t in raw_tokens:
        t = re.sub(r"^[A-Za-z /&-]{2,30}:\s*", "", t).strip()
        if t and len(t) <= 40:
            tokens.append(t)
    counts = Counter(t.lower() for t in tokens)
    first_seen = {}
    for t in tokens:
        first_seen.setdefault(t.lower(), t)
    duplicates = [
        {"skill": first_seen[key], "count": count}
        for key, count in counts.items()
        if count > 1
    ]
    duplicates.sort(key=lambda d: -d["count"])
    return duplicates

def find_overused_skills(text, min_count=4):
    """Flags skills mentioned an unusually high number of times across the
    whole resume — often a sign of deliberate keyword stuffing to game ATS
    keyword matching, which reads as low-effort (and sometimes backfires
    with smarter ATS parsers) rather than helping."""
    occurrences = keyword_processor.extract_keywords(text)
    counts = Counter(occurrences)
    overused = [
        {"skill": skill, "count": count}
        for skill, count in counts.items()
        if count >= min_count
    ]
    overused.sort(key=lambda d: -d["count"])
    return overused

def find_duplicate_sentences(text, min_words=5, max_results=8):
    """Flags sentences (or sentence-like lines) that appear more than once,
    verbatim, anywhere in the resume — a common tell of a resume stitched
    together from a template or from multiple older drafts."""
    raw_lines = [l.strip() for l in text.splitlines() if l.strip()]
    candidates = []
    for line in raw_lines:
        clean = re.sub(r"^[•\-\*\u2022➤◆■●▪]\s*", "", line).strip()
        for sent in SENTENCE_SPLIT_RE.split(clean):
            sent = sent.strip()
            if sent:
                candidates.append(sent)
    counts = Counter()
    first_seen = {}
    for s in candidates:
        if len(s.split()) < min_words:
            continue
        key = re.sub(r"[^\w\s]", "", s.lower())
        key = re.sub(r"\s+", " ", key).strip()
        if not key:
            continue
        counts[key] += 1
        first_seen.setdefault(key, s)
    duplicates = [
        {"sentence": first_seen[key], "count": count}
        for key, count in counts.items()
        if count > 1
    ]
    duplicates.sort(key=lambda d: -d["count"])
    return duplicates[:max_results]

def _normalize_bullet_for_dedup(bullet):
    # Numbers are swapped for a placeholder so two bullets that only differ
    # by a metric (e.g. "grew revenue by 12%" vs "...by 30%") still register
    # as the same underlying claim, copy-pasted across roles.
    b = re.sub(r"\d+", "#", bullet.lower())
    b = re.sub(r"[^\w\s#]", "", b)
    return re.sub(r"\s+", " ", b).strip()

def find_duplicate_bullets(text, similarity_threshold=0.85, max_pairs=8, max_compare=120):
    """Flags bullet points duplicated verbatim across different jobs or
    projects, plus near-identical bullets (same claim reworded only
    slightly) — both are just as noticeable to a recruiter skimming for
    distinct, specific accomplishments per role."""
    bullets = extract_bullet_lines(text)
    if len(bullets) < 2:
        return {"exact": [], "near": []}

    counts = Counter()
    first_seen = {}
    for b in bullets:
        key = _normalize_bullet_for_dedup(b)
        if not key:
            continue
        counts[key] += 1
        first_seen.setdefault(key, b)
    exact = [
        {"bullet": first_seen[key], "count": count}
        for key, count in counts.items()
        if count > 1
    ]
    exact.sort(key=lambda d: -d["count"])

    # Near-duplicates: pairwise comparison over distinct bullets, bounded so
    # this stays manageable even on a long resume (O(n^2) on max_compare).
    distinct = list(dict.fromkeys(bullets))[:max_compare]
    near = []
    for i in range(len(distinct)):
        ka = _normalize_bullet_for_dedup(distinct[i])
        for j in range(i + 1, len(distinct)):
            kb = _normalize_bullet_for_dedup(distinct[j])
            if not ka or not kb or ka == kb:
                continue  # exact duplicates already captured above
            ratio = difflib.SequenceMatcher(None, ka, kb).ratio()
            if ratio >= similarity_threshold:
                near.append({
                    "bullet_a": distinct[i],
                    "bullet_b": distinct[j],
                    "similarity": round(ratio * 100),
                })
    near.sort(key=lambda d: -d["similarity"])
    return {"exact": exact, "near": near[:max_pairs]}

def detect_duplicate_content(text, chunks=None):
    """Top-level duplicate/redundancy scan surfaced in the 'Duplicate
    Content Detection' report card: repeated skills, repeated sentences,
    and repeated (or near-identical) bullet points."""
    if chunks is None:
        chunks = split_sections(text)
    duplicate_skills = find_duplicate_skills(chunks)
    overused_skills = find_overused_skills(text)
    duplicate_sentences = find_duplicate_sentences(text)
    bullets = find_duplicate_bullets(text)
    total_issues = (
        len(duplicate_skills)
        + len(overused_skills)
        + len(duplicate_sentences)
        + len(bullets["exact"])
        + len(bullets["near"])
    )
    return {
        "duplicate_skills": duplicate_skills,
        "overused_skills": overused_skills,
        "duplicate_sentences": duplicate_sentences,
        "duplicate_bullets": bullets["exact"],
        "near_duplicate_bullets": bullets["near"],
        "total_issues": total_issues,
        "has_duplicates": total_issues > 0,
    }